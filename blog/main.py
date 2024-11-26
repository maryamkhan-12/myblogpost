import io, os, re, json, shutil, requests
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from tqdm import tqdm
from PIL import Image
from pprint import pprint
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from typing import List, Dict
from pydantic import BaseModel
from docx.shared import Inches
from docx.shared import RGBColor
from docx.oxml import OxmlElement
from langchain_groq import ChatGroq
from fastapi import FastAPI, HTTPException ,Query
from fastapi.responses import FileResponse
from langchain_core.prompts import PromptTemplate
app = FastAPI()
class ImagePromptRequest(BaseModel):
    blog_post_content: str
    previous_image_prompts: str
llm = ChatGroq(model="llama-3.2-90b-text-preview", temperature=0.3, max_tokens=1024, api_key='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' )
class blog_request(BaseModel):
    TypeOf : str
    target_audience: str
    tone: str
    point_of_view: str
    target_country: str
    keywords: List[str]
def fetch_google_results(keywords: List[str], target_country: str) -> List[str]:
    username = 'marcosmcdonnell@gmail.com'
    password = 'Webscraper1'
    all_results_dict = {}
    for keyword in keywords:
        payload = {'source': 'google_search','query': keyword, 'domain': 'com', 'geo_location': target_country, 'locale': 'en-us', 'parse': True, 'start_page': 1, 'pages': 5,  'limit': 10, }
    try:
        response = requests.post('https://realtime.oxylabs.io/v1/queries', auth=(username, password), json=payload )
        response.raise_for_status()
        all_results_dict[keyword] = response.json()
    except requests.RequestException as e:
        raise HTTPException(status_code=500, detail=f"Error for '{keyword}': {str(e)}")
    formatted_results = {keyword: {'results': [{ 'pos': organic.get('pos'), 'url': organic.get('url'),'title': organic.get('title')}
                for result in all_results_dict[keyword].get('results', [])
                for organic in result.get('content', {}).get('results', {}).get('organic', [])   ]        }    }
    return formatted_results

def generate_blog_title(keywords: List[str], search_results: List[str], category: str,blog_request: blog_request) -> str:
    prompt_template = """
    You are an expert content creator and SEO strategist. Your task is to craft a single, SEO-optimized, and reader-focused title for a blog post using the provided keywords. 

    Blog Post Details:
    - **Category**: {category}
    - **Keywords**: {keywords}
    - **Type**: {Type}

    Instructions:
    1. Use the given keywords naturally and effectively within the title.
    2. Ensure the title is concise (preferably under 60 characters) while retaining clarity and relevance.
    3. Make it catchy and engaging to attract readers’ attention.
    4. Reflect the blog's type and purpose (e.g., listicle, guide, how-to, etc.) in the title, catering specifically to readers interested in {category}.
    5. Maintain a positive, inviting tone that aligns with the topic.
    6. Avoid generic or vague phrases; ensure the title is specific and impactful.

    Output:
    - Provide a single SEO-friendly title only, without additional explanations or formatting.
"""

    prompt = prompt_template.format(
        Type=blog_request.TypeOf,
        category=category,
        keywords=", ".join(keywords),
        search_results="\n".join(search_results)
    )
    response = llm.invoke(prompt)
    return response.content

def generate_blog_subheadings(title: str, search_results : list, seleted_catagory:str, blog_request: blog_request) -> List[str]:
    prompt_template = """
    You are a skilled content strategist and SEO expert tasked with creating compelling and SEO-optimized subheadings for a blog post. These subheadings should enhance readability, engage the target audience, and align with the blog’s title and focus.

    Blog Post Details:
    - **Title**: {title}
    - **Category**: {seleted_catagory}
    - **SEO Keywords**: {search_results}
    - **Target Audience**: {target_audience}
    - **Tone**: {tone}
    - **Point of View**: {point_of_view}
    - **Target Country**: {target_country}

    Instructions:
    1. Generate 2 subheadings that comprehensively address important aspects of the topic.
    2. Ensure each subheading incorporates relevant keywords and resonates with the blog's tone and target audience.
    3. Use concise, clear, and engaging language that encourages readers to continue exploring the blog.
    4. If applicable, include tips, actionable insights, or region-specific details to add value.
    5. Maintain a logical flow between subheadings to create a seamless reading experience.
    6. Write only subheadings, and do not include any additional text or formatting in the output.

    Based on this input, suggest SEO-friendly subheadings for the blog post.
"""

    prompt = prompt_template.format(
        title=title, seleted_catagory=seleted_catagory,search_results=", ".join(search_results), target_audience=blog_request.target_audience, 
        tone=blog_request.tone,point_of_view=blog_request.point_of_view,target_country=blog_request.target_country 
    )
    response = llm.invoke(prompt)
    suggested_subheadings = response.content.split("\n")
    return [subheading.strip() for subheading in suggested_subheadings if subheading.strip()]
    
def BlogPostPromptSingleSubheading(title: str, current_subheading: str, blog_request: blog_request, search_results: List[str], category: str, previous_content: str) -> str:
    prompt_template = """
    You are an expert content creator and language model specializing in crafting professional and engaging blog posts. 
    Your goal is to write a well-structured, SEO-optimized, and captivating section under the given subheading, 
    tailored to the target audience and aligned with the overall blog theme.

    Blog Post Details:
    - **Category**: {category}
    - **Title**: {title}
    - **Target Audience**: {target_audience}
    - **Tone**: {tone}
    - **Point of View**: {point_of_view}
    - **Target Country**: {target_country}
    
    Previous Content (for reference and context):
    {previous_content}
    
    **Subheading**: {current_subheading}

    Instructions:
    1. Write an engaging, unique, and factual section for the given subheading, ensuring it aligns seamlessly with the preceding content.
    2. Optimize the section for SEO by naturally integrating the provided keywords. Do not overuse them; maintain readability and flow.
    3. Use a tone that resonates with the target audience (e.g., reassuring, authoritative, or conversational) and fits the blog’s theme.
    4. Incorporate relevant research findings, statistics, expert quotes, or actionable advice to enrich the content and make it credible.
    5. Include practical tips, relatable examples, or insights that address the audience’s needs, questions, or challenges.
    6. Use smooth transitions to maintain consistency and guide readers into the subsequent sections effortlessly.
    7. Avoid irrelevant details, filler content, or generic phrases. Keep every sentence valuable and impactful.
    8. Use numbers for lists or tips (e.g., "1.", "2.", "3.") for clarity and organization. Do not use ** for lists or tips.
    9. For headings and subheadings, use ** (e.g., **Subheading**) to maintain consistent formatting.

    Additional Notes:
    - Each section should be concise yet comprehensive (maximum 2-3 paragraphs per subheading).
    - Focus on writing for readers first, with SEO considerations seamlessly integrated.
    - Avoid adding conclusions, references, or FAQs in the content.
    - Ensure the writing naturally leads into the next subheading.

    Now, based on this input, draft a compelling and SEO-friendly section for the given subheading.
"""

    prompt = prompt_template.format(title=title, category=category,target_audience=blog_request.target_audience,tone=blog_request.tone,keywords=", ".join(search_results),point_of_view=blog_request.point_of_view,target_country=blog_request.target_country,previous_content=previous_content,current_subheading=current_subheading )
    response = llm.invoke(prompt)
    content=response.content
    return content
    
def format_content(document, content: str):
    subheading_pattern = r"\*\*(.*?)\*\*"
    bullet_point_pattern = r"^\s*•\s*\*\*\s*(.*?)\s*\*\*"
    lines = content.split("\n")
    for line in lines:
        if re.match(subheading_pattern, line):
            subheading_text = re.sub(r"\*\*", "", line).strip()
            document.add_heading(subheading_text, level=2)
        elif re.match(bullet_point_pattern, line):
            bullet_text = re.sub(r"^\s*•\s*\*\*", "", line).strip()
            p = document.add_paragraph(style='List Bullet')
            run = p.add_run(bullet_text)
            run.bold = True
        else:
            p = document.add_paragraph(line.strip())
            p.alignment = 3  # Justify alignment
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.color.rgb = RGBColor(0, 0, 0)  # Set text color to black
    
def generate_image_prompt(content: str, previous_prompts: str) -> str:
    prompt_template = """
You are a creative assistant tasked with generating visually stunning, realistic image prompts for a child-care blog. Each prompt should be specific, emotionally engaging, and professionally crafted to enhance the blog post.

Blog Post Draft:
{blog_post_content}

Previous Image Prompts:
{previous_image_prompts}

Instructions:
- Identify impactful themes or scenes from the blog content to transform into clear and visually captivating image prompts.
- Ensure prompts are concise and include:
  - Specific subject/scene: Include details relevant to child-care topics, such as age-appropriate toys, family interactions, or daily routines.
  - Composition and action: Describe the arrangement of elements, depth, and any storytelling moments.
  - Emotion and style: Capture a specific feeling or artistic inspiration (e.g., candid, editorial, or minimalist).
  - Lighting and color: Suggest lighting techniques (e.g., natural golden-hour light) and color palettes (e.g., soft pastels or warm tones).
  - Camera and lens settings: Recommend professional equipment (e.g., Canon EOS R5, 50mm lens) to match the visual tone.
  - Artistic enhancements: Add textures, angles, or subtle details to make the scene more engaging.
  - Aspect ratio: Suggest proportions (e.g., --ar 16:9 for banners or --ar 4:5 for social media).

Examples:
1. A cozy nursery at the end of a playful day, with soft pastel-colored walls, scattered toys on a colorful rug, and sunlight fading through the curtains. A half-read children’s book rests open on the chair, evoking warmth and care. Shot at golden hour with a Canon EOS 5D Mark IV and a 24-70mm lens. Soft bokeh effect for added depth. --ar 16:9 --style raw
2. A toddler bundled in a bright, cozy winter coat, playing in the snow while building a small snowman, surrounded by frosty trees. Soft morning sunlight creates a dreamy glow, highlighting the child’s joyful expression. Captured with a Canon EOS R5 using a 70-200mm lens for soft background blur. --ar 3:2 --style raw
3. A minimalist flat-lay composition of baby toys, with muted pastel tones, soft lighting, and a clean vignette for a modern aesthetic. Shot with a Canon EOS 90D and a macro lens to enhance texture and detail. --ar 1:1 --style raw
---
"""


    prompt = prompt_template.format(blog_post_content=content,previous_image_prompts=previous_prompts)
    response = llm.invoke(prompt)
    return response.content  # Extract prompts
    
def generate_image(prompt: str):
    # Get the API token from environment variables
    #API_TOKEN = os.getenv("API_TOKEN")
    API_TOKEN = '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'

    
    if not API_TOKEN:
        raise ValueError("API_TOKEN is not set. Please check your secrets configuration.")

    API_URL = "https://api-inference.huggingface.co/models/black-forest-labs/FLUX.1-dev"
    headers = {"Authorization": f"Bearer {API_TOKEN}"}
    payload = {"inputs": prompt}
    response = requests.post(API_URL, headers=headers, json=payload)
    if response.status_code == 200:
        image_bytes = response.content
        image = Image.open(io.BytesIO(image_bytes))
        return image
    else:
        print("Failed to generate image.")
        return None

def selected_category(category: dict, search_results: list) -> str:
    prompt_template = """
    Based on the given search results, select the most appropriate category for the blog post.
    Available Categories: {categories}
    Search Results: 
    {search_results}
    Carefully analyze the keywords and context in the search results to choose the best category. 
    Please respond only with the most relevant category name.
    """
    prompt = prompt_template.format(categories=", ".join(category.keys()), search_results="\n".join(search_results))
    response = llm.invoke(prompt)
    return response.content.strip()  # Extract the selected category
    
def fetch_google_results_for_site(keywords: List[str]) -> List[Dict[str, int]]:
    username = 'yousaf_PaOrl'
    password = 'Black_hawk=2'
    query_string = "+".join(keywords)
    search_url = f"https://www.google.com/search?q=site:marcusmcdonnell.com+{query_string}"
    payload = {
        'source': 'google',
        'url': search_url,
        'parse': True  # Enabling parsed response to get structured data
    }
    try:
        response = requests.post(
            'https://realtime.oxylabs.io/v1/queries',
            auth=(username, password),
            json=payload
        )
        response.raise_for_status()
        full_response = response.json()
        filtered_results = []
        if full_response.get('results'):
            for result in full_response['results']:
                organic_results = result.get('content', {}).get('results', {}).get('organic', [])
                if isinstance(organic_results, list):
                    filtered_results.extend(
                        {"title": item.get("title"), "url": item.get("url"), "pos": item.get("pos")}
                        for item in organic_results
                        if "title" in item and "url" in item and "pos" in item                    )
                else:
                    print("Expected 'organic' results to be a list but found something else.")
        else:
            print("No 'results' key found in the response.")
        return filtered_results
    except requests.RequestException as e:
        print(f"Error fetching results: {e}")
        return []

def decide_to_generate_image(content: str, i:int) -> bool:
    max=3
    prompt_template = '''
    You are an advanced language model tasked with deciding if an image should be generated based on the provided blog post. Analyze the blog content and respond with "Yes" or "No" only. Generate an image only if the content is rich, high-quality, and would benefit from it. Generate a maximum of 3 images: if {i} > {max}, respond with "No".
    Blog post:
    {blog_post}
    Output:
    Yes or No
    '''
    prompt=prompt_template.format(blog_post=content,i=i,max=max)
    response = llm.invoke(prompt)
    should_generate_image = response.content
    return should_generate_image
        
def generate_linkages(blog_post: str, search_results: list, keywords: List[str]) -> dict:
    Internal_search_results = fetch_google_results_for_site(keywords)
    prompt_template = """
    Based on the given blog post and search results, generate relevant external and internal links.
    
    Blog Post:
    {blog_post}
    
    Use the top 3 search results for external link suggestions, considering their relevance and quality. The links should be clickable hyperlinks.
    Also, suggest internal links that might help the reader based on the blog post's content. Do not include placeholder statements like 'no links found.'
    
    External Links:
    Provide a list of up to 3 high-quality external links with a brief description of each link's content and its relevance to the blog post. Ensure all links are clickable.
    
    Internal Links:
    Suggest up to 3 internal links based on the blog post's content. Provide a brief explanation of how each internal link connects to the blog post.
    
    External Links Results: 
    {search_results}
    
    Internal Links Results:
    {Internal_search_results}
    
    Output:
    External Links: 
    - [Link Text](URL): Brief explanation of relevance.
    
    Internal Links:
    - [Link Text](URL): Brief explanation of relevance.
    """

    prompt = prompt_template.format(blog_post=blog_post, search_results=search_results, Internal_search_results=Internal_search_results)
    response = llm.invoke(prompt)
    result = response.content.strip()  # Adjust based on LLM output structure
    
    return result
    
@app.post("/generate_blog/", response_model=dict)
def create_blog_pipeline(blog_request: blog_request):
    try:
        word_file_path = "/tmp/Generated_Blog_Post.docx"
        for file_path in [word_file_path]:
            if os.path.exists(file_path):
                os.remove(file_path)
                
        category = {"Parenting Stages":["Baby & Toddler Years","Preschool & Early Childhood","Big Kids (6–12 Years)","Tweens & Teens","Newborn Care","Parenting After Baby #2 (or #3!)"],"Everyday Life with Kids":["Daily Routines & Organization","Mealtime & Nutrition","Playtime & Activities","Sleep Schedules & Tips","Family Time Ideas","Special Occasions & Holidays"],"Self-Care for Moms":["Health & Wellness","Mental Health & Stress Relief","Beauty & Self-Care Tips","Hobbies & “Me Time”","Personal Growth & Goal Setting"],"Parenting Tips & Tricks":["Time-Saving Hacks","Budgeting for Families","Quick Cleaning Tips","Home Organization with Kids","School & Homework Help","Tech Tools for Parenting"],"Mom Life (Humor & Reality)":["Honest Mom Moments","Laughs & Parenting Memes","Confessions & Fails","Real Life, Real Moms","Quotes & Relatable Stories"],"Parenting Styles & Philosophies":["Gentle Parenting & Positive Discipline","Attachment Parenting","Raising Independent Kids","Balancing Structure & Freedom","Parenting Trends & Research"],"Relationships & Family Dynamics":["Co-Parenting & Communication","Sibling Relationships","Family Bonding Activities","Grandparents & Extended Family","Blended Families & Step-Parenting"],"Learning & Development":["Early Childhood Education","Fun Learning at Home","Language & Social Skills","Emotional Development","Physical & Motor Skills"],"Health & Wellness":["Child Nutrition & Health","Exercise & Play for Kids","Sleep Health","Pediatric Check-Ups","Common Illnesses & Remedies","Childproofing & Safety"],"Mompreneurs & Working Moms":["Balancing Work & Family","Remote Work Tips","Side Hustles & Passions","Time Management for Busy Moms","Self-Care for Working Moms"],"School & Education":["Preschool & Daycare Choices","School Readiness & Transitions","Homework & Study Skills","Navigating School Friendships","Involvement in School Activities"],"Lifestyle & Home":["Home Décor for Families","Sustainable & Eco-Friendly Choices","Family Finances & Budgeting","Travel & Family Adventures","Pets & Kids"],"Parenting Challenges":["Tantrums & Discipline","Social Media & Screen Time","Bullying & Peer Pressure","Dealing with Picky Eaters","Navigating Kids’ Fears"],"Creative & Fun Ideas":["DIY Projects for Kids","Kid-Friendly Crafts","Fun Recipes & Snacks","Family Games & Activities","Fun Celebrations & Birthdays"],"Modern Parenting Topics":["Raising Kids in a Digital World","Cultural & Diversity Awareness","Gender-Neutral Parenting","Parenting and Social Media"],"The Wild World of Parenting":["Tiny Bosses: Life with Toddlers","Kid Logic: Decoding the Mind of a Child","Growing Up Fast: Navigating the Tween Years"],"The Mom Zone":["Surviving the Madness, One Coffee at a Time","Keeping It Real: The Beautiful Mess of Mom Life","Dear Diary: Honest Mom Confessions"],"Mastering the Art of Family Chaos":["Organized Chaos: Tips for Running a Household","Barely Hanging On: Parenting Hacks for the Real World","Kid-Proof Your Life (If That’s Even Possible)"],"Mom Self-Care, Simplified":["Time for You: Self-Care for Busy Moms","Staying Sane (Mostly) with Self-Care on a Budget","Love Yourself: From Self-Care to Self-Love"],"Making Memories, Keeping Your Sanity":["Everyday Magic: Fun Family Traditions","Making the Ordinary Extraordinary","The Cool Mom’s Guide to Family Fun"],"Mom Hacks & Life-Saving Tricks":["Shortcuts for Sanity: Clever Parenting Hacks","The No-Fuss Guide to Getting Stuff Done","Mom Brain Solutions: Little Tricks for Big Wins"],"When Kids Are…Kids!":["Real Talk: Tantrums, Tears & Tiny Attitudes","Kid Quirks: The Weird, Wonderful World of Children","Mini People, Mega Emotions"],"Relationships and Realities":["It Takes Two: Parenting Together (Even When You Don’t Agree)","Keeping Love Alive Amid the Chaos","Keeping the Family Peace, One Day at a Time"],"The School Scene":["Homework Without the Headache","From Preschool to Preteen Drama: Surviving School Years","Winning at School (Even If They Don’t Love It)"],"Digital World for Digital Kids":["Screen Time vs. Play Time: Finding the Balance","Raising Tech-Savvy Kids in a Tech-Obsessed World","Social Media & Selfies: Teaching Digital Smarts"],"Raising the Next Generation":["The Kindness Project: Raising Empathetic Kids","How to Raise Future World-Changers","The Power of Yes and No: Teaching Choices"],"Healthier, Happier Families":["Making Meals Easy & Fun (Yes, Really!)","Health Hacks for Kids Who Hate Veggies","Small Habits for Big Health Wins"],"The Organized Chaos Hub":["Declutter Like a Pro (Yes, Even with Kids)","Home Hacks for the Ultimate Kid-Friendly Space","Mastering the Family Schedule"],"Funny Mom Survival Kit":["Parenting Memes You’ll Feel in Your Soul","Surviving Kids’ Parties with Style","Confessions of a Bedtime Warrior"],"Big Dreams & Little Goals":["Goal-Getting for Moms Who Do It All","Dare to Dream Big (Even If You’re Tired)","Mom Goals: From ‘Just Survive’ to ‘Thrive’"],"For the Love of Learning":["Learning Through Play: Fun Ideas for Little Learners","Home Learning Hacks for Smart Kids","Raising Curious Kids: Sparking Little Imaginations"],"Tales from the Trenches":["Stories from the Wild World of Parenting","Lessons Learned from the Chaos","Hilarious Mom Stories You’ll Never Believe"],"Adventures Big and Small":["Tiny Adventures: Fun for Kids of All Ages","Family Vacations & Kid-Friendly Getaways","Staycations That Feel Like the Real Deal"],"The Support Network":["For the Love of Moms: Support & Community","Village of Moms: Finding Your Support Circle","Surviving & Thriving Together"],"Creative Kids Zone":["Arts & Crafts that Won’t Break the Bank","Imagination Station: Encouraging Creative Play","Rainy Day Fun: Indoor Ideas for Any Weather"]}

        output_dir = "/tmp/pic"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)

        print('SEO Searching')
        search_results = fetch_google_results( blog_request.keywords,  blog_request.target_country        )
        print('Selecting Category for blog post')
        selected_cat = selected_category(category, search_results)
        print('Generating Title for blog post')
        previous_image_prompts = ''
        blog_content = ""
        document = Document()
        title = generate_blog_title(
            blog_request.keywords, search_results, selected_cat, blog_request        )
        document.add_heading(title, 0)
        image_prompt = generate_image_prompt(title, previous_image_prompts)
        previous_image_prompts += image_prompt + " , "
        image = generate_image(image_prompt)
        
        if image:
            title_image_path = "/tmp/pic/image.png"
            image.save(title_image_path)
            document.add_picture(title_image_path, width=Inches(6), height=Inches(6))
        else:
            print("Title image generation failed.")
        print('Generating Subheadings for blog post')
        
        subheadings = generate_blog_subheadings(title, search_results, selected_cat, blog_request )
        
        for i, subheading in enumerate(tqdm(subheadings, desc="Processing subheadings")):
            content = BlogPostPromptSingleSubheading(                
                title, subheading, blog_request, search_results, selected_cat, blog_content            
            )
            
            blog_content += f"\n\n{subheading}\n{content}"
            
            format_content(document, content)
            
            if decide_to_generate_image(content, i):
                image_prompt = generate_image_prompt(content, previous_image_prompts)
                previous_image_prompts += image_prompt + " , "
                image = generate_image(image_prompt)
                if image:
                    subheading_image_path = f"/tmp/pic/image_{i}.png"
                    image.save(subheading_image_path)
                    document.add_picture(subheading_image_path, width=Inches(6), height=Inches(6))
                else:
                    print(f"Image generation failed for subheading: {subheading}")
        
        # Generate linkages
        raw_linkages = generate_linkages(blog_content, search_results, blog_request.keywords)
        formatted_linkages = format_linkages(raw_linkages)

        # Add linkages to the document
        document.add_heading("Relevant Links", level=2)
        document.add_paragraph(formatted_linkages)

        document.save(word_file_path)
        shutil.rmtree("/tmp/pic/", ignore_errors=True)
        return {"docx_path": word_file_path}
    except Exception as e:
        print(f"An error occurred: {e}")
        return {"error": str(e)}

# Update to include format_linkages
def format_linkages(linkages: str) -> str:
    """
    Format the linkages into a readable bullet-point list.
    """
    formatted_linkages = "\n".join([f"- {line.strip()}" for line in linkages.split("\n") if line.strip()])
    return formatted_linkages

@app.get("/download/")
def download_file(file_path: str = Query(...)):
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type='application/octet-stream', filename=os.path.basename(file_path))
    else:
        return {"error": "File not found"}

@app.get("/")
async def root():
    return {"message": "API is up and running!"}